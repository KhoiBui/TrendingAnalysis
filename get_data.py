""" Extract the findings table from final CAPA report. """

import re
from docx import Document
from openpyxl import load_workbook


class GetData(object):
    """ Read the final CAPA and extract info. """

    def __init__(self, document, workbook, worksheet):
        self.document = Document(document)
        self.workbook = load_workbook(workbook)
        self.worksheet = self.workbook.get_sheet_by_name(worksheet)
        self.table_data = []
        self.data_read = []
        self.project_info = {}
        self.table = None
        self.leads = ['Adam', 'Monika', 'Jeff', 'Mario']
        self.findings = ['Process Area', 'Goal', 'Practice', 'Description', 'Rating']

    def process_document(self):
        """ Process the document. """
        self.find_table()
        self.read_doc()
        if self.table is None:
            # no table found
            print("#####    Not able to find \"Detail of Findings\" table.   #####")
            print("##### Possible that project does not have any findings. #####")
            return
        print(self.project_info)
        self.read_table_data(self.table)

    def find_table(self):
        """ Locate the detailed findings table. """
        tables = self.document.tables
        header = []
        for table in tables:
            for row in table.rows:
                header[:] = []
                for cell in row.cells:
                    for para in cell.paragraphs:
                        header.append(para.text.strip(' '))
                # new versions of final CAPA's keep project information in a table
                if 'Project Information' in header:
                    self.read_new_format(table)
                # check if elements in findings is also in header
                cond = len(header) == 5 and header[4] == 'Rating'
                if cond or [x for x in self.findings for y in header if x in y] == self.findings:
                    self.table = table
                    return

    def read_doc(self):
        """ Read document and put info into list. """
        self.data_read[:] = []
        for para in self.document.paragraphs:
            text = para.text
            # skip blank lines
            if text.strip():
                # remove duplicated spaces
                text = ' '.join(text.split())
                # for older versions of final CAPA's
                self.fill_project_info(text, new_format=False)
                self.data_read.append(text)

        # Constant in old & new report format
        # Batch/Project name
        # Lead(s)'s name
        # Reported date
        for i in range(0, len(self.data_read)):
            if next((x for x in self.leads if x in self.data_read[i]), None):
                self.project_info.update({'Project Name': self.data_read[i - 1]})
                self.project_info.update({'Lead(s)': self.data_read[i]})
                self.project_info.update({'Date Reported': self.data_read[i + 1]})
                break

    def fill_project_info(self, line_read, new_format):
        """ Get general information about the project from doc. """
        if new_format:
            """ list(set()) converts list into a set to remove duplicates, but
                does not preserve order. dict key must always be 1st element
                in line_read else the project_info dict will not be updated
                properly. The code below gets rid of duplicated values while
                preserving the order of the elements in line_read. Code was
                taken from a StackOverflow thread about the same issue. """
            tmp_list = set()    # new empty set
            tmp_add = tmp_list.add    # built-in method 'add' of set object
            temp = [x for x in line_read if not (x in tmp_list or tmp_add(x))]
            line_read = temp
        else:
            line_read = line_read.split(':', 1)
        line_read[0] = re.sub('[- ]', '', line_read[0])
        key_name = line_read[0].lower()

        if 'sapid' in key_name:
            self.project_info.update({'SAP ID': line_read[1].strip(' ')})
        if 'golive' in key_name:
            self.project_info.update({'Go Live Date': line_read[1]})
        if 'customer' in key_name:
            site_name = re.sub('State|Lottery', '', line_read[1])
            site_name = site_name.strip(' ')
            self.project_info.update({'Site': site_name})

    def read_new_format(self, table):
        """ Read the project information in new format of CAPA report. """
        data = []
        index = 0
        for row in table.rows:
            data.append([])
            for cell in row.cells:
                text = ''
                for para in cell.paragraphs:
                    text += para.text.strip(' ')
                data[index].append(text)
            self.fill_project_info(data[index], new_format=True)
            index += 1

    def read_table_data(self, table):
        """ Read info in specified table. """
        data = []
        index = 0
        for row in table.rows:
            data.append([])
            for cell in row.cells:
                text_data = ''
                for para in cell.paragraphs:
                    text_data += para.text.strip(' ')
                data[index].append(text_data)
            index += 1

        # trim unneeded rows in old & new reports
        if all('CAPA' in x for x in data[0]):
            self.table_data = data[2:]
        else:
            self.table_data = data[1:]
        # trim end of list
        self.table_data = [row[:5] for row in self.table_data]

    def get_table_data(self):
        """ Return data in table. """
        return self.table_data

    def get_project_info(self):
        """ Return information about the project. """
        return self.project_info

    def get_doc_data(self):
        """ Return information about the project. """
        return self.data_read

    def get_worksheet(self):
        """ Return name of the worksheet. """
        return self.worksheet

    def get_workbook(self):
        """ Return name of the workbook. """
        return self.workbook

    def get_document(self):
        """ Return name of the document. """
        return self.document
