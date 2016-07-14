""" Extract the findings table from final CAPA report. """

import re
from docx import Document
from openpyxl import load_workbook

class ExtractTable(object):
    """ help. """

    table_data = []
    data_read = []
    project_info = {}
    table = ''
    findings = ['Process Area', 'Goal', 'Practice', 'Description', 'Rating']

    def __init__(self, workbook, worksheet, document):
        self.workbook = load_workbook(workbook, read_only=False)
        self.worksheet = self.workbook.get_sheet_by_name(worksheet)
        self.document = Document(document)

    def process_document(self):
        """ Process the document. """
        self.read_doc()
        self.find_table()
        self.read_table_data(self.table)
        self.project_info.update({'Project Name':self.data_read[2]})
        self.project_info.update({'Lead(s)':self.data_read[3]})
        self.project_info.update({'Date Reported':self.data_read[4]})

    def find_table(self):
        """ help. """
        if self.findings is None:
            raise ValueError('Row header is invalid')

        tables = self.document.tables
        header = []
        for table in tables:
            header_row = table.rows[0]
            header[:] = []
            for cell in header_row.cells:
                for para in cell.paragraphs:
                    header.append(para.text.strip(' '))
            if header == self.findings:
                self.table = table
                return

    def fill_project_info(self, line_read):
        """ help. """
        line_read = line_read.split(':', 1)
        line_read[0] = re.sub('[- ]', '', line_read[0])
        key_name = line_read[0].lower()

        if 'sap' in key_name:
            self.project_info.update({'SAP ID':line_read[1]})
        elif 'golive' in key_name:
            self.project_info.update({'Go Live Data':line_read[1]})

    def read_doc(self):
        """ help. """
        for para in self.document.paragraphs:
            text = para.text
            # skip blank lines
            if text is not '':
                # remove duplicated spaces
                text = ' '.join(text.split())
                self.fill_project_info(text)
                self.data_read.append(text)

    def read_table_data(self, table):
        """ help. """
        data = []
        index = -1
        for row in table.rows:
            data.append([])
            index += 1
            for cell in row.cells:
                for para in cell.paragraphs:
                    data[index].append(para.text.strip(' '))

        # don't need header row anymore
        self.table_data = data[1:]

    def get_table_data(self):
        """ Return data in table. """
        return self.table_data

    def get_project_info(self):
        """ Return information about the project. """
        return self.project_info

    def get_doc_data(self):
        """ Return information about the project. """
        return self.data_read
